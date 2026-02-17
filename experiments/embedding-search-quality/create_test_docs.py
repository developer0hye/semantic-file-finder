"""
검색 품질 테스트용 다양한 주제의 문서 생성.

각 문서는 서로 다른 도메인/주제를 다뤄서,
자연어 검색 시 올바른 문서가 상위에 랭크되는지 확인할 수 있게 한다.
"""

from pathlib import Path
from docx import Document
from pptx import Presentation
from pptx.util import Inches, Pt

OUTPUT_DIR = Path(__file__).parent / "test_files"


def create_docs():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    docs = {
        # 1. 재무 보고서
        "finance_q3_revenue_report.docx": {
            "title": "2024년 3분기 매출 실적 보고서",
            "paragraphs": [
                "2024년 3분기 매출은 전년 동기 대비 15% 증가한 198억원을 기록하였습니다.",
                "주요 성장 동인은 클라우드 서비스 부문으로, 해당 부문 매출이 45% 급증했습니다.",
                "영업이익률은 21.3%로 전분기 대비 2.1%p 개선되었으며, "
                "순이익은 41억원을 달성했습니다.",
                "4분기에는 신규 엔터프라이즈 고객 확보와 글로벌 시장 확대를 통해 "
                "연간 매출 목표 달성이 가능할 것으로 전망됩니다.",
            ],
        },
        # 2. 채용 공고
        "hr_backend_developer_jd.docx": {
            "title": "백엔드 개발자 채용 공고",
            "paragraphs": [
                "당사는 시리즈 B 스타트업으로, 백엔드 개발자를 모집합니다.",
                "주요 업무: REST API 설계 및 개발, 마이크로서비스 아키텍처 구축, "
                "데이터베이스 설계 및 최적화, CI/CD 파이프라인 관리.",
                "자격 요건: Python 또는 Go 3년 이상 경력, PostgreSQL 및 Redis 경험, "
                "AWS 또는 GCP 클라우드 인프라 경험, Docker/Kubernetes 실무 경험.",
                "우대 사항: 대규모 트래픽 처리 경험, 오픈소스 기여 이력.",
                "연봉: 6,000만원 ~ 9,000만원 (경력에 따라 협의).",
                "근무지: 서울 강남구, 하이브리드 근무 가능.",
            ],
        },
        # 3. 프로젝트 제안서
        "project_ai_chatbot_proposal.docx": {
            "title": "AI 고객 상담 챗봇 도입 제안서",
            "paragraphs": [
                "본 제안서는 고객 상담 업무 효율화를 위한 AI 챗봇 시스템 도입을 제안합니다.",
                "현황: 월평균 상담 건수 12,000건, 평균 응답 시간 4.5분, 고객 만족도 72%.",
                "도입 후 기대 효과: 단순 문의의 70%를 자동 응답 처리, "
                "평균 응답 시간 30초 이내로 단축, 고객 만족도 90% 이상 달성.",
                "기술 스택: GPT-4 기반 자연어 처리, RAG(검색 증강 생성) 파이프라인, "
                "Elasticsearch 기반 FAQ 검색, React 기반 채팅 UI.",
                "예상 비용: 초기 구축 1.5억원, 월 운영비 500만원.",
                "구축 기간: 3개월 (요구사항 분석 1개월, 개발 1.5개월, 테스트 0.5개월).",
            ],
        },
        # 4. 계약서
        "legal_service_contract.docx": {
            "title": "소프트웨어 개발 용역 계약서",
            "paragraphs": [
                "계약 당사자: (갑) 주식회사 테크코프, (을) 주식회사 데브파트너.",
                "계약 기간: 2024년 7월 1일부터 2024년 12월 31일까지 (6개월).",
                "계약 금액: 총 3억원 (부가가치세 별도). 착수금 30%, 중간금 40%, 잔금 30%.",
                "개발 범위: ERP 시스템 고도화, 모바일 앱 개발, API 연동 개발.",
                "지식재산권: 개발 결과물의 소유권은 갑에게 귀속되며, "
                "을은 유사 프로젝트에 핵심 기술을 재사용할 수 있는 권리를 보유한다.",
                "하자보수: 검수 완료 후 12개월간 무상 하자보수를 제공한다.",
                "위약벌: 일방의 귀책사유로 계약을 해지할 경우 계약금액의 10%를 위약벌로 지급한다.",
            ],
        },
        # 5. 회의록
        "meeting_product_launch_20241015.docx": {
            "title": "신제품 출시 회의록 (2024.10.15)",
            "paragraphs": [
                "참석자: 김철수(PM), 이영희(마케팅), 박민수(개발), 정수진(디자인), 최동욱(QA).",
                "안건 1: 출시 일정 확인 — 11월 15일 정식 출시 확정. "
                "베타 테스트는 10월 25일부터 2주간 진행.",
                "안건 2: 마케팅 전략 — 출시 전 2주간 티저 캠페인 진행. "
                "인플루언서 10명에게 사전 체험 제공. 출시 당일 프레스 릴리즈 배포.",
                "안건 3: 기술 이슈 — 결제 모듈 PG사 연동 테스트 완료 필요. "
                "성능 테스트에서 동시 접속 1만명 기준 응답시간 2초 초과 이슈 발견. "
                "박민수 담당으로 10월 20일까지 해결 예정.",
                "안건 4: 디자인 — 앱 아이콘 최종 시안 3종 중 B안으로 확정. "
                "다크모드는 2차 업데이트에서 지원.",
                "다음 회의: 10월 22일 오후 2시.",
            ],
        },
        # 6. 기술 문서
        "tech_api_documentation.docx": {
            "title": "사용자 인증 API 기술 문서 v2.1",
            "paragraphs": [
                "본 문서는 OAuth 2.0 기반 사용자 인증 API의 기술 명세를 기술합니다.",
                "Base URL: https://api.example.com/v2/auth",
                "POST /login — 사용자 로그인. Request Body: { email, password }. "
                "Response: { access_token, refresh_token, expires_in }.",
                "POST /refresh — 토큰 갱신. Header: Authorization: Bearer {refresh_token}. "
                "Response: { access_token, expires_in }.",
                "POST /logout — 로그아웃. 현재 세션의 모든 토큰을 무효화.",
                "Rate Limit: 로그인 시도 분당 5회, 일반 API 호출 분당 100회.",
                "에러 코드: 401 Unauthorized, 429 Too Many Requests, 500 Internal Server Error.",
                "access_token 유효기간: 1시간. refresh_token 유효기간: 30일.",
            ],
        },
        # 7. 마케팅 전략
        "marketing_2025_strategy.docx": {
            "title": "2025년 마케팅 전략 계획서",
            "paragraphs": [
                "2025년 마케팅 예산: 총 15억원 (전년 대비 20% 증가).",
                "핵심 목표: 브랜드 인지도 50% → 70% 향상, MAU 30만 → 50만 달성, "
                "고객 획득 비용(CAC) 15% 절감.",
                "채널별 예산 배분: 디지털 광고 40%, 콘텐츠 마케팅 25%, "
                "오프라인 이벤트 15%, PR 10%, 인플루언서 마케팅 10%.",
                "상반기 전략: SEO 강화 및 블로그 콘텐츠 주 3회 발행. "
                "구글/메타 광고 캠페인 A/B 테스트 진행.",
                "하반기 전략: 연말 프로모션 및 파트너십 마케팅 집중. "
                "오프라인 컨퍼런스 2회 개최.",
            ],
        },
        # 8. 연구 보고서
        "research_battery_technology.docx": {
            "title": "차세대 리튬-황 배터리 기술 동향 보고서",
            "paragraphs": [
                "리튬-황(Li-S) 배터리는 이론 에너지 밀도 2,600 Wh/kg으로 "
                "기존 리튬이온 배터리(250-300 Wh/kg)의 약 10배에 달한다.",
                "주요 과제: 폴리설파이드 셔틀 효과로 인한 용량 감소, "
                "충방전 사이클 수명 500회 미만, 황 양극의 낮은 전도성.",
                "최신 연구 동향: 탄소 나노튜브 기반 양극 구조로 사이클 수명 1,000회 달성 (MIT, 2024). "
                "고체 전해질 적용으로 셔틀 효과 80% 감소 (KAIST, 2024).",
                "상용화 전망: 2027년 드론/UAV용 시장 진입 예상. "
                "2030년 전기차용 본격 상용화 목표.",
                "투자 동향: 글로벌 Li-S 배터리 스타트업 투자 누적 약 42억 달러 (2024년 기준).",
            ],
        },
        # 9. 출장 보고서
        "trip_report_tokyo_2024.docx": {
            "title": "도쿄 출장 보고서 (2024.09.23 ~ 09.25)",
            "paragraphs": [
                "출장자: 김민재 (해외사업팀). 출장지: 일본 도쿄.",
                "목적: 일본 현지 파트너사 NexTech Japan과의 기술 협력 미팅 및 "
                "2025년 공동 사업 계획 논의.",
                "주요 내용: NexTech Japan의 AI 영상 분석 기술과 당사 클라우드 플랫폼의 "
                "통합 솔루션 개발에 합의. POC 프로젝트 1월 착수 예정.",
                "경비: 항공료 85만원, 숙박비 45만원(2박), 식비 15만원, 교통비 8만원, "
                "합계 153만원.",
                "후속 조치: 11월 중 MOU 체결, 12월 기술 교류 워크숍 개최.",
            ],
        },
        # 10. 교육 자료
        "training_git_basics.docx": {
            "title": "Git 기초 교육 자료",
            "paragraphs": [
                "Git은 분산 버전 관리 시스템으로, 소스 코드의 변경 이력을 추적합니다.",
                "기본 명령어: git init (저장소 생성), git add (스테이징), "
                "git commit (커밋), git push (원격 전송), git pull (원격 동기화).",
                "브랜치 전략: main(배포용), develop(개발용), feature/*(기능 개발). "
                "기능 개발 완료 후 Pull Request를 통해 develop에 병합.",
                "커밋 메시지 규칙: 'feat: 새 기능', 'fix: 버그 수정', "
                "'refactor: 리팩토링', 'docs: 문서 수정'.",
                "충돌 해결: git merge 시 충돌이 발생하면 해당 파일을 수동으로 편집 후 "
                "git add → git commit으로 해결.",
            ],
        },
    }

    for filename, content in docs.items():
        doc = Document()
        doc.add_heading(content["title"], level=1)
        for p in content["paragraphs"]:
            doc.add_paragraph(p)
        path = OUTPUT_DIR / filename
        doc.save(str(path))
        print(f"Created: {path}")

    print(f"\nTotal: {len(docs)} documents in {OUTPUT_DIR}/")


if __name__ == "__main__":
    create_docs()
